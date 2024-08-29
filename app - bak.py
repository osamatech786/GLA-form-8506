import streamlit as st
from streamlit_drawable_canvas import st_canvas
import shutil
import re
from PIL import Image as PILImage
from datetime import datetime, date, timedelta
import time
import smtplib
from email.message import EmailMessage
import os
from docx import Document
from docx.shared import Inches
# from dotenv import load_dotenv



files=list()
# mandatory fields validation
exclude_fields = {}     
mandatory_fields = []


def app():
    st.set_page_config(
        page_title="Prevista - GLA Form",
        page_icon="📝",
        layout="wide",
        initial_sidebar_state="collapsed",
    )

    st.image('header/header-GLA.png', use_column_width=True)

    global mandatory_fields

    st.title('Welcome')
    st.subheader('Please fill out the following details:')

    title_mr, title_mrs, title_miss, title_ms='','','',''
    title = st.radio(
        "Title",
        ["Mr", "Mrs", "Miss", "Ms"]
    )
    if title == "Mr":
        title_mr = 'X'
    elif title == "Mrs":
        title_mrs = 'X'
    elif title == "Miss":
        title_miss = 'X'
    elif title == "Ms":
        title_ms = 'X'


    first_name = st.text_input('First Name', key="first_name")
    middle_name = st.text_input('Middle Name', key="middle_name")
    family_name = st.text_input('Family Name', key="family_name")
    # mandatory_fields.extend([f'p{i}' for i in range(1, 4)]) 

    # Initialize gender variables
    gender_m, gender_f, other_gender, other_gender_text = '', '', '', ''
    # Radio button for gender selection
    gender = st.radio("Gender", ["M", "F", "Other"])
    # Conditional input for "Other" gender option
    if gender == "M":
        gender_m = 'X'
    elif gender == "F":
        gender_f = 'X'
    elif gender == "Other":
        other_gender =  'X'
        other_gender_text = st.text_input("If Other, please state")
        # mandatory_fields.extend(['p117'])
    
    date_of_birth = st.date_input(
    label="Date of Birth",
    value=datetime(2000, 1, 1),  # Default date
    min_value=date(1900, 1, 1),  # Minimum selectable date
    max_value=date(2025, 12, 31),  # Maximum selectable date
    key="date_input_widget",  # Unique key for the widget
    help="Choose a date",  # Tooltip text
    format='DD/MM/YYYY'
)
    current_age = calculate_age(date_of_birth)
    date_of_birth = date_of_birth.strftime("%d-%m-%Y")
    
    current_age_text='Current Age at Start of Programme: '+ str(current_age)
    st.text(current_age_text)

    ethnicity_options = {
        'White': {
            'English/ Welsh/ Scottish/ N Irish/ British': '31',
            'Irish': '32',
            'Roma, Gypsy or Irish Traveller': '33',
            'Any other white background': '34'
        },
        'Mixed/ Multiple ethnic group': {
            'White and Black Caribbean': '35',
            'White and Black African': '36',
            'White and Asian': '37',
            'Any other mixed/ multiple ethnic background': '38'
        },
        'Asian/ Asian British': {
            'Indian': '39',
            'Pakistani': '40',
            'Bangladeshi': '41',
            'Chinese': '42',
            'Any other Asian background': '43'
        },
        'Black/ African/ Caribbean/ Black British': {
            'African': '44',
            'Caribbean': '45',
            'Any Other Black/ African/ Caribbean background': '46'
        },
        'Other Ethnic Group': {
            'Arab': '47',
            'Any other ethnic group': '98'
        }
    }

    ethnicity_category = st.selectbox('Select Ethnicity Category', list(ethnicity_options.keys()))
    ethnicity = st.selectbox('Select Ethnicity', list(ethnicity_options[ethnicity_category].keys()))
    ethnicity_code = ethnicity_options[ethnicity_category][ethnicity]
    st.write(f'Ethnicity Code: {ethnicity_code}')

    # Initialize ethnicity variables
    ethnicity_vars = {f'ethnicity_{i}': '' for i in range(31, 48)}
    ethnicity_48=''
    # Set the corresponding ethnicity variable to 'X'
    if ethnicity_code == 31:
        ethnicity_vars['ethnicity_31'] = 'X'
    elif ethnicity_code == 32:
        ethnicity_vars['ethnicity_32'] = 'X'
    elif ethnicity_code == 33:
        ethnicity_vars['ethnicity_33'] = 'X'
    elif ethnicity_code == 34:
        ethnicity_vars['ethnicity_34'] = 'X'
    elif ethnicity_code == 35:
        ethnicity_vars['ethnicity_35'] = 'X'
    elif ethnicity_code == 36:
        ethnicity_vars['ethnicity_36'] = 'X'
    elif ethnicity_code == 37:
        ethnicity_vars['ethnicity_37'] = 'X'
    elif ethnicity_code == 38:
        ethnicity_vars['ethnicity_38'] = 'X'
    elif ethnicity_code == 39:
        ethnicity_vars['ethnicity_39'] = 'X'
    elif ethnicity_code == 40:
        ethnicity_vars['ethnicity_40'] = 'X'
    elif ethnicity_code == 41:
        ethnicity_vars['ethnicity_41'] = 'X'
    elif ethnicity_code == 42:
        ethnicity_vars['ethnicity_42'] = 'X'
    elif ethnicity_code == 43:
        ethnicity_vars['ethnicity_43'] = 'X'
    elif ethnicity_code == 44:
        ethnicity_vars['ethnicity_44'] = 'X'
    elif ethnicity_code == 45:
        ethnicity_vars['ethnicity_45'] = 'X'
    elif ethnicity_code == 46:
        ethnicity_vars['ethnicity_46'] = 'X'
    elif ethnicity_code == 47:
        ethnicity_vars['ethnicity_47'] = 'X'
    else:
        ethnicity_48='X'


    national_insurance_number = st.text_input("National Insurance Number")

    house_no_name_street = st.text_input("House No./Name & Street")
    suburb_village = st.text_input("Suburb / Village")
    town_city = st.text_input("Town / City")
    county = st.text_input("County")
    country_of_domicile = st.text_input("Country of Domicile")
    current_postcode = st.text_input("Current Postcode")
    postcode_prior_enrollment = st.text_input("Postcode Prior to Enrolment")
    email_address = st.text_input("Email Address")
    primary_telephone_number = st.text_input("Primary Telephone Number")
    secondary_telephone_number = st.text_input("Secondary Telephone Number")
    next_of_kin = st.text_input("Next of kin/Emergency contact")
    emergency_contact_phone_number = st.text_input("Emergency Contact Phone Number")

    # mandatory_fields.extend([f'p{i}' for i in range(137, 150)])

    # Household Situation Section
    st.header('Household Situation')
    st.subheader('Please select the most relevant options. (Tick ALL relevant boxes)')

    household_options = {
        '1 - No household member in employment with one or more dependent children': 'JH, JH+DC',
        '2 - No household member in employment with no dependent children': 'JH',
        '3 - Participant lives in a single adult household with dependent children': 'SAH+DC',
        '4 - Learner lives in single unemployed adult household with dependent children': 'JH, SAH+DC',
        '99 - None of the above apply': 'N/A'
    }

    # Store household selections
    household_selections = {}
    for option, code in household_options.items():
        household_selections[option] = st.checkbox(option, key=code)

    # Initialize relevant variables with empty string values
    no_member_employed_with_children = ''
    no_member_employed_without_children = ''
    single_adult_household_with_children = ''
    unemployed_single_adult_household = ''
    none_of_the_above = ''

    # Set variables based on selections
    if household_selections.get('1 - No household member in employment with one or more dependent children'):
        no_member_employed_with_children = 'X'
    if household_selections.get('2 - No household member in employment with no dependent children'):
        no_member_employed_without_children = 'X'
    if household_selections.get('3 - Participant lives in a single adult household with dependent children'):
        single_adult_household_with_children = 'X'
    if household_selections.get('4 - Learner lives in single unemployed adult household with dependent children'):
        unemployed_single_adult_household = 'X'
    if household_selections.get('99 - None of the above apply'):
        none_of_the_above = 'X'
        
    # # Display selected household situations
    # st.subheader('Selected Household Situations:')
    # selected_households = [option for option, selected in household_selections.items() if selected]
    # if selected_households:
    #     for selected in selected_households:
    #         st.write(selected)
    # else:
    #     st.write('No options selected.')

    # Check if at least one checkbox is selected
    if any(household_selections.values()):
        household_filled = 'filled'
    else:
        household_filled = ''

    # Extend the mandatory_fields list with the household_filled variable
    # mandatory_fields.extend(['p300'])


    # LLDD, Health Problems, Other Disadvantaged Section
    st.header('LLDD, Health Problems, Other Disadvantaged')

    # Long term disability, health problem, or learning difficulties
    st.subheader('Do you consider yourself to have a long term disability, health problem or any learning difficulties? Choose the correct option. If Yes enter code in Primary LLDD or HP; you can add multiple LLDD or HP but primary must be recorded if Yes selected.')
    disability = st.radio('Choose the correct option:', ['N', 'Y'], index=0)
    # Initialize variables for disability options
    has_disability, no_disability = '', ''
    
    # initilize first to overcome error:
    # Initialize variables for each health problem type
    vision_impairment_primary, vision_impairment_secondary, vision_impairment_tertiary = '-', '-', '-'
    hearing_impairment_primary, hearing_impairment_secondary, hearing_impairment_tertiary = '-', '-', '-'
    mobility_impairment_primary, mobility_impairment_secondary, mobility_impairment_tertiary = '-', '-', '-'
    complex_disabilities_primary, complex_disabilities_secondary, complex_disabilities_tertiary = '-', '-', '-'
    social_emotional_difficulties_primary, social_emotional_difficulties_secondary, social_emotional_difficulties_tertiary = '-', '-', '-'
    mental_health_difficulty_primary, mental_health_difficulty_secondary, mental_health_difficulty_tertiary = '-', '-', '-'
    moderate_learning_difficulty_primary, moderate_learning_difficulty_secondary, moderate_learning_difficulty_tertiary = '-', '-', '-'
    severe_learning_difficulty_primary, severe_learning_difficulty_secondary, severe_learning_difficulty_tertiary = '-', '-', '-'
    dyslexia_primary, dyslexia_secondary, dyslexia_tertiary = '-', '-', '-'
    dyscalculia_primary, dyscalculia_secondary, dyscalculia_tertiary = '-', '-', '-'
    autism_spectrum_primary, autism_spectrum_secondary, autism_spectrum_tertiary = '-', '-', '-'
    aspergers_primary, aspergers_secondary, aspergers_tertiary = '-', '-', '-'
    temporary_disability_primary, temporary_disability_secondary, temporary_disability_tertiary = '-', '-', '-'
    speech_communication_needs_primary, speech_communication_needs_secondary, speech_communication_needs_tertiary = '-', '-', '-'
    physical_disability_primary, physical_disability_secondary, physical_disability_tertiary = '-', '-', '-'
    specific_learning_difficulty_primary, specific_learning_difficulty_secondary, specific_learning_difficulty_tertiary = '-', '-', '-'
    medical_condition_primary, medical_condition_secondary, medical_condition_tertiary = '-', '-', '-'
    other_learning_difficulty_primary, other_learning_difficulty_secondary, other_learning_difficulty_tertiary = '-', '-', '-'
    other_disability_primary, other_disability_secondary, other_disability_tertiary = '-', '-', '-'
    prefer_not_to_say= '-'
    additional_info=''

    # Set variables based on user selection
    if disability == 'Y':
        has_disability, no_disability = 'Y', '-'

        # LLDD or Health Problem Types
        st.subheader('LLDD or Health Problem Type')

        

        # Health problem types data
        data = [
            ('Vision impairment (4)', 'vision_primary', 'vision_secondary', 'vision_tertiary'),
            ('Hearing impairment (5)', 'hearing_primary', 'hearing_secondary', 'hearing_tertiary'),
            ('Disability affecting mobility (6)', 'mobility_primary', 'mobility_secondary', 'mobility_tertiary'),
            ('Profound complex disabilities (7)', 'complex_primary', 'complex_secondary', 'complex_tertiary'),
            ('Social and emotional difficulties (8)', 'social_primary', 'social_secondary', 'social_tertiary'),
            ('Mental health difficulty (9)', 'mental_primary', 'mental_secondary', 'mental_tertiary'),
            ('Moderate learning difficulty (10)', 'moderate_primary', 'moderate_secondary', 'moderate_tertiary'),
            ('Severe learning difficulty (11)', 'severe_primary', 'severe_secondary', 'severe_tertiary'),
            ('Dyslexia (12)', 'dyslexia_primary', 'dyslexia_secondary', 'dyslexia_tertiary'),
            ('Dyscalculia (13)', 'dyscalculia_primary', 'dyscalculia_secondary', 'dyscalculia_tertiary'),
            ('Autism spectrum disorder (14)', 'autism_primary', 'autism_secondary', 'autism_tertiary'),
            ('Asperger\'s syndrome (15)', 'aspergers_primary', 'aspergers_secondary', 'aspergers_tertiary'),
            ('Temporary disability after illness (for example post-viral) or accident (16)', 'temporary_primary', 'temporary_secondary', 'temporary_tertiary'),
            ('Speech, Language and Communication Needs (17)', 'speech_primary', 'speech_secondary', 'speech_tertiary'),
            ('Other physical disability (18)', 'physical_primary', 'physical_secondary', 'physical_tertiary'),
            ('Other specific learning difficulty (e.g. Dyspraxia) (19)', 'specific_primary', 'specific_secondary', 'specific_tertiary'),
            ('Other medical condition (for example epilepsy, asthma, diabetes) (20)', 'medical_primary', 'medical_secondary', 'medical_tertiary'),
            ('Other learning difficulty (90)', 'other_learning_primary', 'other_learning_secondary', 'other_learning_tertiary'),
            ('Other disability (97)', 'other_disability_primary', 'other_disability_secondary', 'other_disability_tertiary'),
            ('Prefer not to say (98)', 'prefer_not_to_say', '', '')
        ]

        # Starting placeholder index
        placeholder_index = 157

        # Create checkboxes and map them to variables explicitly
        for label, primary, secondary, tertiary in data:
            st.write(f'**{label}**')
            
            # Create checkboxes
            primary_checked = st.checkbox('Primary', key=primary)
            secondary_checked = st.checkbox('Secondary', key=secondary) if secondary else False
            tertiary_checked = st.checkbox('Tertiary', key=tertiary) if tertiary else False

            # Set variables based on selections
            if primary_checked:
                if 'vision' in primary:
                    vision_impairment_primary = 'X'
                elif 'hearing' in primary:
                    hearing_impairment_primary = 'X'
                elif 'mobility' in primary:
                    mobility_impairment_primary = 'X'
                elif 'complex' in primary:
                    complex_disabilities_primary = 'X'
                elif 'social' in primary:
                    social_emotional_difficulties_primary = 'X'
                elif 'mental' in primary:
                    mental_health_difficulty_primary = 'X'
                elif 'moderate' in primary:
                    moderate_learning_difficulty_primary = 'X'
                elif 'severe' in primary:
                    severe_learning_difficulty_primary = 'X'
                elif 'dyslexia' in primary:
                    dyslexia_primary = 'X'
                elif 'dyscalculia' in primary:
                    dyscalculia_primary = 'X'
                elif 'autism' in primary:
                    autism_spectrum_primary = 'X'
                elif 'asperger' in primary:
                    aspergers_primary = 'X'
                elif 'temporary' in primary:
                    temporary_disability_primary = 'X'
                elif 'speech' in primary:
                    speech_communication_needs_primary = 'X'
                elif 'physical' in primary:
                    physical_disability_primary = 'X'
                elif 'specific' in primary:
                    specific_learning_difficulty_primary = 'X'
                elif 'medical' in primary:
                    medical_condition_primary = 'X'
                elif 'other_learning' in primary:
                    other_learning_difficulty_primary = 'X'
                elif 'other_disability' in primary:
                    other_disability_primary = 'X'
                elif 'prefer_not' in primary:
                            prefer_not_to_say = 'X'

            if secondary_checked:
                if 'vision' in secondary:
                    vision_impairment_secondary = 'X'
                elif 'hearing' in secondary:
                    hearing_impairment_secondary = 'X'
                elif 'mobility' in secondary:
                    mobility_impairment_secondary = 'X'
                elif 'complex' in secondary:
                    complex_disabilities_secondary = 'X'
                elif 'social' in secondary:
                    social_emotional_difficulties_secondary = 'X'
                elif 'mental' in secondary:
                    mental_health_difficulty_secondary = 'X'
                elif 'moderate' in secondary:
                    moderate_learning_difficulty_secondary = 'X'
                elif 'severe' in secondary:
                    severe_learning_difficulty_secondary = 'X'
                elif 'dyslexia' in secondary:
                    dyslexia_secondary = 'X'
                elif 'dyscalculia' in secondary:
                    dyscalculia_secondary = 'X'
                elif 'autism' in secondary:
                    autism_spectrum_secondary = 'X'
                elif 'asperger' in secondary:
                    aspergers_secondary = 'X'
                elif 'temporary' in secondary:
                    temporary_disability_secondary = 'X'
                elif 'speech' in secondary:
                    speech_communication_needs_secondary = 'X'
                elif 'physical' in secondary:
                    physical_disability_secondary = 'X'
                elif 'specific' in secondary:
                    specific_learning_difficulty_secondary = 'X'
                elif 'medical' in secondary:
                    medical_condition_secondary = 'X'
                elif 'other_learning' in secondary:
                    other_learning_difficulty_secondary = 'X'
                elif 'other_disability' in secondary:
                    other_disability_secondary = 'X'

            if tertiary_checked:
                if 'vision' in tertiary:
                    vision_impairment_tertiary = 'X'
                elif 'hearing' in tertiary:
                    hearing_impairment_tertiary = 'X'
                elif 'mobility' in tertiary:
                    mobility_impairment_tertiary = 'X'
                elif 'complex' in tertiary:
                    complex_disabilities_tertiary = 'X'
                elif 'social' in tertiary:
                    social_emotional_difficulties_tertiary = 'X'
                elif 'mental' in tertiary:
                    mental_health_difficulty_tertiary = 'X'
                elif 'moderate' in tertiary:
                    moderate_learning_difficulty_tertiary = 'X'
                elif 'severe' in tertiary:
                    severe_learning_difficulty_tertiary = 'X'
                elif 'dyslexia' in tertiary:
                    dyslexia_tertiary = 'X'
                elif 'dyscalculia' in tertiary:
                    dyscalculia_tertiary = 'X'
                elif 'autism' in tertiary:
                    autism_spectrum_tertiary = 'X'
                elif 'asperger' in tertiary:
                    aspergers_tertiary = 'X'
                elif 'temporary' in tertiary:
                    temporary_disability_tertiary = 'X'
                elif 'speech' in tertiary:
                    speech_communication_needs_tertiary = 'X'
                elif 'physical' in tertiary:
                    physical_disability_tertiary = 'X'
                elif 'specific' in tertiary:
                    specific_learning_difficulty_tertiary = 'X'
                elif 'medical' in tertiary:
                    medical_condition_tertiary = 'X'
                elif 'other_learning' in tertiary:
                    other_learning_difficulty_tertiary = 'X'
                elif 'other_disability' in tertiary:
                    other_disability_tertiary = 'X'


        # Additional information that may impact learning
        additional_info = st.text_area('Is there any other additional information that may impact on your ability to learn?')


    else:
        has_disability, no_disability = '-', 'N'

    
    # Other disadvantaged sections
    st.subheader('Other disadvantaged - Ex Offender?')
    ex_offender = st.radio('', ['Y', 'N', 'Choose not to say'], key='ex_offender')
    # Initialize ex_offender variables
    ex_offender_y, ex_offender_n, ex_offender_choose_not_to_say = '', '', ''
    # Conditional input for ex_offender option
    if ex_offender == "Y":
        ex_offender_y = 'Y'
    elif ex_offender == "N":
        ex_offender_n = 'N'
    elif ex_offender == "Choose not to say":
        ex_offender_choose_not_to_say = 'Choose not to say'
    

    st.subheader('Other disadvantaged - Homeless?')
    homeless = st.radio('', ['Y', 'N', 'Choose not to say'], key='homeless')
    # Initialize homeless variables
    homeless_y, homeless_n, homeless_choose_not_to_say = '', '', ''
    # Conditional input for homeless option
    if homeless == "Y":
        homeless_y = 'Y'
    elif homeless == "N":
        homeless_n = 'N'
    elif homeless == "Choose not to say":
        homeless_choose_not_to_say = 'Choose not to say'


    # Referral Source Section
    st.header('Referral Source')
    # Creating columns for referral source options
    col1, col2, col3, col4 = st.columns(4)

    # Initialize referral source variables
    internally_sourced, recommendation, event, self_referral, family_friends = '', '', '', '', ''
    other, website, promotional_material, jobcentre_plus = '', '', '', ''
    event_specify, other_specify = '', ''

    # Adding checkboxes for each referral source option
    with col1:
        internally_sourced = st.checkbox('Internally sourced', key='internally_sourced')
        recommendation = st.checkbox('Recommendation')
        promotional_material = st.checkbox('Promotional material')
    with col2:
        self_referral = st.checkbox('Self Referral')
        family_friends = st.checkbox('Family/ Friends')
        event = st.checkbox('Event (please specify)')
    with col3:
        website = st.checkbox('Website', key='website')
        jobcentre_plus = st.checkbox('JobCentre Plus')
        other = st.checkbox('Other (please specify)')
    # Text inputs for 'Event (please specify)' and 'Other (please specify)' if checked
    if event:
        event_specify = st.text_input('Please specify the event', key='event_specify')
    if other:
        other_specify = st.text_input('Please specify other source', key='other_specify')

    specify_refereel = st.text_input("Please let us know the organization or advisor who referred you to our program, or indicate where you found out about this opportunity. If it was through a job center, please specify its location.", key="specify_refereel")

    # Setting 'X' for chosen options
    internally_sourced_val = 'X' if internally_sourced else ''
    recommendation_val = 'X' if recommendation else ''
    event_val = event_specify if event else ''
    self_referral_val = 'X' if self_referral else ''
    family_friends_val = 'X' if family_friends else ''
    other_val = other_specify if other else ''
    website_val = 'X' if website else ''
    promotional_material_val = 'X' if promotional_material else ''
    jobcentre_plus_val = 'X' if jobcentre_plus else ''
    


    # mandatory validation
    referrall=''
    if (internally_sourced_val == 'X' or 
        recommendation_val == 'X' or 
        event_val == 'X' or 
        self_referral_val == 'X' or 
        family_friends_val == 'X' or 
        other_val == 'X' or 
        website_val == 'X' or 
        promotional_material_val == 'X' or
        jobcentre_plus_val == 'X' and
        len(specify_refereel)>0):
        referrall = 'filled'
        st.write('LENGTH:', len(specify_refereel))
    mandatory_fields.extend(['p304'])
   
    # Employment and Monitoring Information Section
    st.header('Employment and Monitoring Information')

    # Initialize employment status variables
    unemployed_val, economically_inactive_val, employed_val = '', '', ''

    # Participant Employment Status
    st.subheader('Participant Employment Status')
    employment_status = st.radio(
        "Select your employment status:",
        [
            "Unemployed (looking for work and available to start work) -> go to section A",
            "Economically Inactive (not looking for work and not available to start work) -> Go to section B",
            "Employed (including self-employed) -> go to section C"
        ]
    )

    # Setting 'X' for chosen employment status
    if employment_status == "Unemployed (looking for work and available to start work) -> go to section A":
        unemployed_val = 'X'
    elif employment_status == "Economically Inactive (not looking for work and not available to start work) -> Go to section B":
        economically_inactive_val = 'X'
    elif employment_status == "Employed (including self-employed) -> go to section C":
        employed_val = 'X'

    up_to_12_months_val, twelve_months_or_longer_val = '-', '-'
    # Section A - Unemployment details
    if "Unemployed" in employment_status:
        st.subheader('Section A - Unemployment details')
        st.text("Where a participant’s employment status is long-term unemployed proof of both unemployment and the length of unemployment must be obtained.")
        
        unemployment_duration = st.radio("If you are not working, how long have you been without work?", ["Up to 12 months", "12 months or longer"])
        # Initialize unemployment duration variables
        # Setting 'X' for chosen unemployment duration
        if unemployment_duration == "Up to 12 months":
            up_to_12_months_val = 'X'
        elif unemployment_duration == "12 months or longer":
            twelve_months_or_longer_val = 'X'
                
        # Evidence of Unemployment Status Section
        st.write("Evidence of unemployment status (for more information look Start-Eligibility Evidence list tab)")
        unemployment_evidence = st.selectbox(
            "Select evidence type:",
            [
                "A Letter or Document from JCP or DWP",
                "A written referral from a careers service",
                "Third Party Verification or Referral form",
                "Other (please specify)"
            ]
        )

        # Initialize unemployment evidence variables
        jcp_dwp_val, careers_service_val, third_party_val, other_evidence_val = '-', '-', '-', '-'

        # Setting 'X' for chosen evidence type
        if unemployment_evidence == "A Letter or Document from JCP or DWP":
            jcp_dwp_val = 'X'
            uploaded_file = st.file_uploader("Upload Document from JCP or DWP", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
        elif unemployment_evidence == "A written referral from a careers service":
            careers_service_val = 'X'
            uploaded_file = st.file_uploader("Upload written referral from a careers service", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
        elif unemployment_evidence == "Third Party Verification or Referral form":
            third_party_val = 'X'
            uploaded_file = st.file_uploader("Upload Third Party Verification or Referral form", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
        elif unemployment_evidence == "Other (please specify)":
            other_evidence_val = st.text_input("Please specify other evidence")    

        

    # Initialize economically inactive variables
    inactive_status_val, inactive_evidence_type_val, inactive_evidence_date_val = '-', '-', '-'
    
    # Section B - Economically Inactive details
    if "Economically Inactive" in employment_status:
        st.subheader('Section B - Economically Inactive details')
        
        
        inactive_status = st.radio(
            "The Participant is not employed and does not claim benefits at the time of the enrolment.",
            ["Y", "N"]
        )

        # Setting 'X' for chosen inactive status
        inactive_status_val = 'Yes' if inactive_status == "Y" else 'No'

        inactive_evidence_type_val = st.text_input("Type of evidence for Economically Inactive Status including self-declaration statement.")
        inactive_evidence_date_val = st.date_input("Date of issue of evidence", format='DD/MM/YYYY')
        inactive_evidence_date_val = inactive_evidence_date_val.strftime("%d-%m-%Y")


    # Initialize employment detail variables
    employer_name_val, employer_address_1_val, employer_address_2_val = '', '', ''
    employer_address_3_val, employer_postcode_val, employer_contact_name_val = '', '', ''
    employer_contact_position_val, employer_contact_email_val, employer_contact_phone_val = '', '', ''
    employer_edrs_number_val, living_wage_val, employment_hours_val_0, employment_hours_val_6 = '', '', '', ''
    claiming_benefits_val, sole_claimant_val, benefits_list_val = '', '', ''
    other_benefit_val, benefit_claim_date_val = '', ''
    
    # Initialize variables for benefits
    universal_credit_val = ''
    job_seekers_allowance_val = ''
    employment_support_allowance_val = ''
    incapacity_benefit_val = ''
    personal_independence_payment_val = ''

    # Section C - Employment details
    if "Employed" in employment_status:
        st.subheader('Section C - Employment details')
        
        

        employer_name_val = st.text_input("Employer Name")
        employer_address_1_val = st.text_input("Employer Address 1")
        employer_address_2_val = st.text_input("Employer Address 2")
        employer_address_3_val = st.text_input("Employer Address 3")
        employer_postcode_val = st.text_input("Employer Postcode")
        employer_contact_name_val = st.text_input("Main Employer Contact Name")
        employer_contact_position_val = st.text_input("Contact Position")
        employer_contact_email_val = st.text_input("Contact Email Address")
        employer_contact_phone_val = st.text_input("Contact Telephone Number")
        employer_edrs_number_val = st.text_input("Employer EDRS number")

        living_wage = st.radio("Do you earn more than the National Living Wage of £20,319.00 pa (£10.42ph for 37.5 hrs pw)?", ["Y", "N"])
        living_wage_val = 'Y' if living_wage == "Y" else 'N'

        employment_hours = st.radio("Employment Hours (place an X in the applicable box)", ["0-15 hrs per week", "16+ hrs per week"])
        employment_hours_val_0 = 'X' if employment_hours == "0-15 hrs per week" else '-' 
        employment_hours_val_6 = 'X' if employment_hours == "16+ hrs per week" else '-' 


    st.header("Benefits Detail")
    claiming_benefits = st.radio("Are you claiming any benefits? If so, please describe below what they are.", ["N", "Y"])
    claiming_benefits_val = 'Y' if claiming_benefits == "Y" else 'N'

    
    if claiming_benefits == "Y":
        sole_claimant = st.radio("Are you the sole claimant of the benefit?", ["Y", "N"])
        sole_claimant_val = 'Y' if sole_claimant == "Y" else 'N'


        # Benefits List Section
        benefits_list = st.multiselect(
            "Select the benefits you are claiming:",
            [
                "Universal Credit (UC)",
                "Job Seekers Allowance (JSA)",
                "Employment and Support Allowance (ESA)",
                "Incapacity Benefit (or any other sickness related benefit)",
                "Personal Independence Payment (PIP)",
                "Other - please state"
            ]
        )

        # Update the respective variables based on user selections
        if "Universal Credit (UC)" in benefits_list:
            universal_credit_val = 'X'
        if "Job Seekers Allowance (JSA)" in benefits_list:
            job_seekers_allowance_val = 'X'
        if "Employment and Support Allowance (ESA)" in benefits_list:
            employment_support_allowance_val = 'X'
        if "Incapacity Benefit (or any other sickness related benefit)" in benefits_list:
            incapacity_benefit_val = 'X'
        if "Personal Independence Payment (PIP)" in benefits_list:
            personal_independence_payment_val = 'X'

        # Handle "Other - please state" input
        other_benefit_val = ''
        if "Other - please state" in benefits_list:
            other_benefit_val = st.text_input("Please state other benefit")

        # Input for the date of claim
        benefit_claim_date_val = st.date_input("From what date has the above claim been in effect?", format='DD/MM/YYYY')
        benefit_claim_date_val = benefit_claim_date_val.strftime("%d-%m-%Y")



    # # Detailed Learning Plan Section
    # st.header('Detailed Learning Plan')

    # qualification_reference = st.text_input("Qualification Reference")
    # region_of_work = st.text_input("Region of Work")
    # qualification_course_title = st.text_input("Qualification/Course/Unit Title/Non-Regulated activity")
    # awarding_body = st.text_input("Awarding Body")

    # GLH = st.text_input("GLH")

    # benefit_to_you = st.text_area("What is the benefit to you in completing this learning aim? Please be specific")

    # planned_start_date = st.date_input("Planned Start Date")
    # planned_end_date = st.date_input("Planned End Date", help="Note: Actual End Date to be recorded on 'Outcome and Progression' form at the end of the programme")
    # delivery_postcode = st.text_input("Delivery Postcode")
    # date_of_first_review = st.date_input("Date of first review")

    # st.subheader("Progression - Indicate below the progression planned for this participant when they have completed all training")
    # progression_options = st.multiselect(
    #     "Select progression options:",
    #     [
    #         "Progression within Work",
    #         "Progression into Further Education or Training",
    #         "Progression to Apprenticeship",
    #         "Progression into employment"
    #     ]
    # )

    # progression_aim = st.text_area("Please detail your progression aim")

    # st.subheader("Social Outcomes - How do you rate yourself now out of 5 for the below. 5= Great 1= Poor")

    # health_and_well_being = st.slider("Health and well being", 1, 5, 1)
    # social_integration = st.slider("Social integration", 1, 5, 1)
    # learner_self_efficacy = st.slider("Learner self-efficacy", 1, 5, 1)
    # participation_in_volunteering = st.slider("Participation in volunteering", 1, 5, 1)








    # st.header('Eligibility Check')

    # st.text("""
    #     Evidence CANNOT be accepted that has been entered at a later date than Actual End Date of the start aim.
    #     Evidence must be present for ALL 4 (EO1,2,3,4) of the below eligibility checks. Original documentation must have been witnessed by the Provider and preferably copies made as evidence in case of future audits.
    #     For list of ALL acceptable supporting documents check 'Start-Eligibility Evidence list'
    #     """)

    # st.text("""
    #     UK, EEA Nationals and Non-EEA Nationals

    #     EEA Countries (as of 27/01/2021): 
    #     Austria, Belgium, Bulgaria, Croatia, Republic of Cyprus, Czech Republic, Denmark, Estonia, Finland, France, Germany, Greece, Hungary, Ireland, Italy, Latvia, Lithuania, Luxembourg, Malta, Netherlands, Poland, Portugal, Romania, Slovakia, Slovenia, Spain, Sweden, Iceland, Liechtenstein, Norway.

    #     Switzerland is not an EU or EEA member but is part of the single market. This means Swiss nationals have the same rights to live and work in the UK as other EEA nationals.

    #     “Irish citizens in the UK hold a unique status under each country’s national law. You do not need permission to enter or remain in the UK, including a visa, any form of residence permit or employment permit”. Quote taken from below link:
    #     https://www.gov.uk/government/publications/common-travel-area-guidance/common-travel-area-guidance

    #     Non-EEA nationals who hold leave to enter or leave to remain with a permission to work (including status under the EUSS where they are an eligible family member of an EEA national) are also eligible for ESF support whilst in the UK.
    #     """)

    st.header('E01: Right to Live and Work in the UK')

    # var initialize
    hold_settled_status, hold_pre_settled_status, hold_leave_to_remain = '-', '-', '-'
    not_nationality, passport_non_eu, letter_uk_immigration, passport_endorsed, identity_card, country_of_issue, id_document_reference_number, e01_date_of_issue, e01_date_of_expiry, e01_additional_notes ='-', '-', '-', '-', '-', '-', '-', '-', '-', '-'

    # Create a radio button for the Yes/No question
    british_or_not = st.radio(
        'Are you a UK OR Irish National OR European Economic Area (EEA) National?',
        ('Yes', 'No')
    )

    nationality='-'
    full_uk_passport, full_eu_passport, national_identity_card = '-', '-', '-'
    if british_or_not == 'Yes':
        nationality = st.text_input('Nationality')
        options = [
            'Full UK Passport',
            'Full EU Member Passport (must be in date - usually 10 years)',
            'National Identity Card (EU)'
        ]
        selected_option_nationality = st.radio("Select the type of document:", options)

        if selected_option_nationality == options[0]:
            full_uk_passport, full_eu_passport, national_identity_card = 'X', '', ''
            st.text('Please upload a copy of your Full UK Passport')
            uploaded_file = st.file_uploader("Upload Full UK Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                files.append(uploaded_file_2)

        elif selected_option_nationality == options[1]:
            full_uk_passport, full_eu_passport, national_identity_card = '', 'X', ''
            st.text('Please upload a copy of your Full EU Member Passport')
            uploaded_file = st.file_uploader("Upload Full EU Member Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                files.append(uploaded_file_2)

        elif selected_option_nationality == options[2]:
            full_uk_passport, full_eu_passport, national_identity_card = '', '', 'X'
            st.text('Please upload a copy of your National Identity Card (EU)')
            uploaded_file = st.file_uploader("Upload National Identity Card (EU)", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                files.append(uploaded_file_2)

        if selected_option_nationality in [options[1], options[2]]:
            st.text(
                'In order to be eligible for ESF funding, EEA Nationals must meet one of the following conditions'
            )
            conditions = [
                'a. Hold settled status granted under the EU Settlement Scheme (EUSS)',
                'b. Hold pre-settled status granted under the European Union Settlement Scheme (EUSS)',
                'c. Hold leave to remain with permission to work granted under the new Points Based Immigration System'
            ]
            settled_status = st.radio("Select your status:", conditions)

            if settled_status == conditions[0]:
                hold_settled_status, hold_pre_settled_status, hold_leave_to_remain = 'X', '', ''
                st.text('Please upload your share code which is accessible from the following link:')
                uploaded_file = st.file_uploader("https://www.gov.uk/check-immigration-status", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    files.append(uploaded_file)
                uploaded_file_3 = st.file_uploader("Optional - Upload Back Side of Document ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_3 is not None:
                    files.append(uploaded_file_3)

            elif settled_status == conditions[1]:
                hold_settled_status, hold_pre_settled_status, hold_leave_to_remain = '', 'X', ''
                st.text('Please upload your share code which is accessible from the following link:')
                uploaded_file = st.file_uploader("https://www.gov.uk/check-immigration-status", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    files.append(uploaded_file)
                uploaded_file_3 = st.file_uploader("Optional - Upload Back Side of Document  ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_3 is not None:
                    files.append(uploaded_file_3)

            elif settled_status == conditions[2]:
                hold_settled_status, hold_pre_settled_status, hold_leave_to_remain = '', '', 'X'
                st.text('Please upload your share code which is accessible from the following link:')
                uploaded_file = st.file_uploader("https://www.gov.uk/check-immigration-status", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file is not None:
                    files.append(uploaded_file)
                uploaded_file_3 = st.file_uploader("Optional - Upload Back Side of Document   ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
                if uploaded_file_3 is not None:
                    files.append(uploaded_file_3)

    else:
        not_nationality = st.text_input('Nationality ')
        passport_non_eu_checked = st.checkbox(
            'Passport from non-EU member state (must be in date) AND any of the below a, b, or c'
        )
        if passport_non_eu_checked:
            passport_non_eu = 'X'
            st.text('Please upload a copy of your non-EU Passport')
            uploaded_file = st.file_uploader("Upload Non-EU Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_2 = st.file_uploader("Optional - Upload Back Side of Document", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_2 is not None:
                files.append(uploaded_file_2)
        else:
            passport_non_eu = ''

        document_options = [
            "a. Letter from the UK Immigration and Nationality Directorate granting indefinite leave to remain (settled status)",
            "b. Passport either endorsed 'indefinite leave to remain' – (settled status) or includes work or residency permits or visa stamps (unexpired) and all related conditions met; add details below",
            "c. Some non-EEA nationals have an Identity Card (Biometric Permit) issued by the Home Office in place of a visa, confirming the participant’s right to stay, work or study in the UK – these cards are acceptable"
        ]

        document_type = st.radio("Select the type of document:", document_options)
        letter_uk_immigration, passport_endorsed, identity_card = '', '', ''

        if document_type == document_options[0]:
            letter_uk_immigration, passport_endorsed, identity_card = 'X', '', ''
            st.text('Please upload your Letter from the UK Immigration and Nationality Directorate')
            uploaded_file = st.file_uploader("Upload Letter from UK Immigration and Nationality Directorate", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_4 = st.file_uploader("Optional - Upload Back Side of Document ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_4 is not None:
                files.append(uploaded_file_4)

        elif document_type == document_options[1]:
            letter_uk_immigration, passport_endorsed, identity_card = '', 'X', ''
            st.text('Please upload your endorsed passport')
            uploaded_file = st.file_uploader("Upload Endorsed Passport", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_4 = st.file_uploader("Optional - Upload Back Side of Document  ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_4 is not None:
                files.append(uploaded_file_4)

        elif document_type == document_options[2]:
            letter_uk_immigration, passport_endorsed, identity_card = '', '', 'X'
            st.text('Please upload your Identity Card (Biometric Permit)')
            uploaded_file = st.file_uploader("Upload Identity Card (Biometric Permit)", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file is not None:
                files.append(uploaded_file)
            uploaded_file_4 = st.file_uploader("Optional - Upload Back Side of Document   ", type=['docx', 'pdf', 'jpg', 'jpeg', 'png'])
            if uploaded_file_4 is not None:
                files.append(uploaded_file_4)

        country_of_issue = st.text_input('Country of issue')
        id_document_reference_number = st.text_input('ID Document Reference Number')

        e01_date_of_issue = st.date_input(
            label="Date of Issue",
            value=datetime(2000, 1, 1),  # Default date
            min_value=date(1900, 1, 1),  # Minimum selectable date
            max_value=date(2025, 12, 31),  # Maximum selectable date
            help="Choose a date",  # Tooltip text
            format='DD/MM/YYYY'
        )
        e01_date_of_issue = e01_date_of_issue.strftime("%d-%m-%Y")

        e01_date_of_expiry = st.date_input(
            label="Date of Expiry",
            value=datetime(2000, 1, 1),  # Default date
            min_value=date(1900, 1, 1),  # Minimum selectable date
            max_value=date(2050, 12, 31),  # Maximum selectable date
            help="Choose a date",  # Tooltip text
            format='DD/MM/YYYY'
        )
        e01_date_of_expiry = e01_date_of_expiry.strftime("%d-%m-%Y")

        e01_additional_notes = st.text_area('Additional Notes',
                                            'Use this space for additional notes where relevant (type of Visa, restrictions, expiry etc.)')
        


    st.header('E02: Proof of Age (* all documents must be in date and if a letter is used, it must be within the last 3 months)')

    full_passport_eu = add_checkbox_with_upload('Full Passport (EU Member State)', 'full_passport_eu')
    national_id_card_eu = add_checkbox_with_upload('National ID Card (EU)', 'national_id_card_eu')
    firearms_certificate = add_checkbox_with_upload('Firearms Certificate/Shotgun Licence', 'firearms_certificate')
    birth_adoption_certificate = add_checkbox_with_upload('Birth/Adoption Certificate', 'birth_adoption_certificate')
    e02_drivers_license = add_checkbox_with_upload('Drivers Licence (photo card)', 'e02_drivers_license')
    edu_institution_letter = add_checkbox_with_upload('Letter from Educational Institution* (showing DOB)', 'edu_institution_letter')
    e02_employment_contract = add_checkbox_with_upload('Employment Contract/Pay Slip (showing DOB)', 'e02_employment_contract')
    state_benefits_letter = add_checkbox_with_upload('State Benefits Letter* (showing DOB)', 'state_benefits_letter')
    pension_statement = add_checkbox_with_upload('Pension Statement* (showing DOB)', 'pension_statement')
    northern_ireland_voters_card = add_checkbox_with_upload('Northern Ireland voters card', 'northern_ireland_voters_card')
    
    e02_other_evidence_text=''
    e02_other_evidence_text = st.text_input('Other Evidence: Please state type')

    # Validation for the last 3 months
    current_date = date.today()
    three_months_ago = current_date - timedelta(days=90)

    e02_date_of_issue = st.date_input(
        label="Date of Issue of evidence",
        value=date.today(),  # Default date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date(2025, 12, 31),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        format='DD/MM/YYYY'
    )

    # # Check if the selected date is within the last three months
    # if e02_date_of_issue < three_months_ago:
    #     st.warning("The date of issue is not within the last 3 months. Please select a valid date.")
    #     st.stop()
    # st.success("The date of issue is within the last 3 months.")
    
    e02_date_of_issue = e02_date_of_issue.strftime("%d-%m-%Y")

    # Validation for mandatory field
    documents = [
    full_passport_eu,
    national_id_card_eu,
    firearms_certificate,
    birth_adoption_certificate,
    e02_drivers_license,
    edu_institution_letter,
    e02_employment_contract,
    state_benefits_letter,
    pension_statement,
    northern_ireland_voters_card
    ]

    # Check if at least one of the variables is 'X' or if e02_other_evidence_text is not empty
    if any(doc == 'X' for doc in documents) or e02_other_evidence_text != '':
        e02_filled='Filled'
    else:
        e02_filled=''
    # mandatory_fields.extend(['p301'])
    

    st.header('E03: Proof of Residence (must show the address recorded on ILP) *within the last 3 months')

    e03_drivers_license = add_checkbox_with_upload('Drivers Licence (photo card)', 'e03_drivers_license')
    bank_statement = add_checkbox_with_upload('Bank Statement *', 'bank_statement')
    e03_pension_statement = add_checkbox_with_upload('Pension Statement*', 'e03_pension_statement')
    mortgage_statement = add_checkbox_with_upload('Mortgage Statement*', 'mortgage_statement')
    utility_bill = add_checkbox_with_upload('Utility Bill* (excluding mobile phone)', 'utility_bill')
    council_tax_statement = add_checkbox_with_upload('Council Tax annual statement or monthly bill*', 'council_tax_statement')
    electoral_role_evidence = add_checkbox_with_upload('Electoral Role registration evidence*', 'electoral_role_evidence')
    homeowner_letter = add_checkbox_with_upload('Letter/confirmation from homeowner (family/lodging)', 'homeowner_letter')

    e03_other_evidence_text=''
    e03_other_evidence_text = st.text_input('Other Evidence: Please state type', key='e03_other_evidence')

    # Validation for the last 3 months
    e03_date_of_issue = st.date_input(
        label="Date of Issue evidence",
        value=date.today(),  # Default date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date(2025, 12, 31),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        key='e03_date_of_issue',
        format='DD/MM/YYYY'
    )

    # Check if the selected date is within the last three months
    if e03_date_of_issue < three_months_ago:
        st.warning("The date of issue is not within the last 3 months. Please select a valid date.")
        st.stop()
    st.success("The date of issue is within the last 3 months.")
    e03_date_of_issue = e03_date_of_issue.strftime("%d-%m-%Y")

    # Validation for mandatory field
    documents = [
        e03_drivers_license,
        bank_statement,
        e03_pension_statement,
        mortgage_statement,
        utility_bill,
        council_tax_statement,
        electoral_role_evidence,
        homeowner_letter,
    ]

    # Check if at least one of the variables is 'X' or if e02_other_evidence_text is not empty
    if any(doc == 'X' for doc in documents) or e03_other_evidence_text != '':
        e03_filled='Filled'
    else:
        e03_filled=''
    # mandatory_fields.extend(['p302'])

    st.header('E04: Employment Status (please select one option from below and take a copy)')

    latest_payslip = '-'
    e04_employment_contract = '-'
    confirmation_from_employer = '-'
    redundancy_notice = '-'
    sa302_declaration = '-'
    ni_contributions = '-'
    business_records = '-'
    companies_house_records = '-'
    other_evidence_employed = '-'
    unemployed = '-'

    main_options = [
        'a. Latest Payslip (maximum 3 months prior to start date)',
        'b. Employment Contract',
        'c. Confirmation from the employer that the Participant is currently employed by them which must detail: Participant full name, contracted hours, start date AND date of birth or NINO',
        'd. Redundancy consultation or notice (general notice to group of staff or individual notifications) At risk of Redundancy only',
        'e. Self-employed',
        'f. Other evidence as listed in the \'Start-Eligibility Evidence list\' under Employed section - State below',
        'g. Unemployed (complete the Employment section in ILP form)'
    ]

    selected_main_option = st.radio("Select an employment status or document:", main_options, key="e04_main_option")

    if selected_main_option == main_options[0]:
        latest_payslip = handle_file_upload('Latest Payslip (maximum 3 months prior to start date)', 'latest_payslip')
    elif selected_main_option == main_options[1]:
        e04_employment_contract = handle_file_upload('Employment Contract', 'e04_employment_contract')
    elif selected_main_option == main_options[2]:
        confirmation_from_employer = handle_file_upload('Confirmation from the employer', 'confirmation_from_employer')
    elif selected_main_option == main_options[3]:
        redundancy_notice = handle_file_upload('Redundancy consultation or notice', 'redundancy_notice')
    elif selected_main_option == main_options[4]:
        self_employed_options = [
            "HMRC 'SA302' self-assessment tax declaration, with acknowledgement of receipt (within last 12 months)",
            'Records to show actual payment of Class 2 National Insurance Contributions (within last 12 months)',
            'Business records in the name of the business - evidence that a business has been established and is active / operating (within last 12 months)',
            'If registered as a Limited company: Companies House records / listed as Company Director (within last 12 months)'
        ]
        selected_self_employed_option = st.radio("Select self-employed evidence:", self_employed_options, key="e04_self_employed_option")
        if selected_self_employed_option == self_employed_options[0]:
            sa302_declaration = handle_file_upload("HMRC 'SA302' self-assessment tax declaration", 'sa302_declaration')
        elif selected_self_employed_option == self_employed_options[1]:
            ni_contributions = handle_file_upload('Records of Class 2 National Insurance Contributions', 'ni_contributions')
        elif selected_self_employed_option == self_employed_options[2]:
            business_records = handle_file_upload('Business records', 'business_records')
        elif selected_self_employed_option == self_employed_options[3]:
            companies_house_records = handle_file_upload('Companies House records', 'companies_house_records')
    elif selected_main_option == main_options[5]:
        other_evidence_employed = handle_file_upload("Other evidence as listed in the 'Start-Eligibility Evidence list'", 'other_evidence_employed')
    elif selected_main_option == main_options[6]:
        unemployed = handle_file_upload('Unemployed (complete the Employment section in ILP form)', 'unemployed')

    # Validation for the date of issue
    current_date = date.today()
    three_months_ago = current_date - timedelta(days=90)

    e04_date_of_issue = st.date_input(
        label="Date of Issue of evidence",
        value=date.today(),  # Default date
        min_value=date(1900, 1, 1),  # Minimum selectable date
        max_value=date(2025, 12, 31),  # Maximum selectable date
        help="Choose a date",  # Tooltip text
        key='e04_date_of_issue',
        format='DD/MM/YYYY'
    )

    if e04_date_of_issue < three_months_ago:
        st.warning("The date of issue is not within the last 3 months. Please select a valid date.")
        st.stop()
    st.success("The date of issue is within the last 3 months.")
    e04_date_of_issue = e04_date_of_issue.strftime("%d-%m-%Y")

    st.header('Details of Qualification or Training')
  
    qualification_or_training = st.radio(
    'Are you currently undertaking a qualification or training?',
    ['No', 'Yes'])

    if qualification_or_training=='Yes':
        qualification_or_training_y, qualification_or_training_n = 'Y', '-'

        course_details = st.text_area('Course Details',
                                      'Enter details of the course')
        funding_details = st.text_area(
            'Funding Details', 'Enter details of how the course is funded')
    else:
        qualification_or_training_y, qualification_or_training_n = '-', 'N'
        course_details, funding_details = '', ''
        st.write(
            'You answered "No" to currently undertaking a qualification or training.'
        )

    st.header('Evidenced Qualification Levels')


    st.subheader('Participant self declaration of highest qualification level')
    participant_options = [
        'Below Level 1', 'Level 1', 'Level 2', 'Full Level 2', 'Level 3', 'Full Level 3', 'Level 4',
        'Level 5', 'Level 6', 'Level 7 and above', 'No Qualifications'
    ]


    participant_declaration = st.radio('', participant_options)


    p58 = '-'
    p59 = '-'
    p60 = '-'
    p60z = '-'
    p60a = '-'
    p61 = '-'
    p61z = '-'
    p61a = '-'
    p62 = '-'
    p63 = '-'
    p63z = '-'
    p63a = '-'
    p63b = '-'
    p64 = '-'


    if participant_declaration == participant_options[0]:   #Below Level 1
        p58 = 'X'
    elif participant_declaration == participant_options[1]: #Level 1
        p59 = 'X'
    elif participant_declaration == participant_options[2]: #Level 2
        p60, p60z = 'X', 'X'
    elif participant_declaration == participant_options[3]: #Full Level 2
        p60, p60a = 'X', 'X'
    elif participant_declaration == participant_options[4]: #Level 3
        p61, p61z = 'X', 'X'
    elif participant_declaration == participant_options[5]: #Full Level 3
        p61, p61a = 'X', 'X'
    elif participant_declaration == participant_options[6]: #Level 4
        p62 = 'X'
    elif participant_declaration == participant_options[7]: #Level 5
        p63, p63z = 'X', 'X' 
    elif participant_declaration == participant_options[8]: #Level 6
        p63, p63a = 'X', 'X'
    elif participant_declaration == participant_options[9]: #Level 7 and above
        p63, p63b = 'X', 'X', 'X'
    elif participant_declaration == participant_options[10]: #No Qualifications
        p64 = 'X'
    


    # st.subheader('Training Providers declaration')
    # training_provider_options = [
    #     'Below Level 1', 'Level 1', 'Level 2', 'Level 3', 'Below Level 4',
    #     'Level 5 and above', 'No Qualifications', 'No Personal Learning Record'
    # ]

    # training_provider_declaration = st.radio(
    #     'Please check the PLR and record information about prior attainment level to ensure correct recording of prior attainment, as well as ensuring no duplication of learning aims or units takes place.',
    #     training_provider_options)
    # p65 = '-'
    # p66 = '-'
    # p67 = '-'
    # p68 = '-'
    # p69 = '-'
    # p70 = '-'
    # p71 = '-'
    # p72 = '-'
    # justification='-'


    # if training_provider_declaration == training_provider_options[0]:
    #     p65 = 'X'
    # elif training_provider_declaration == training_provider_options[1]:
    #     p66 = 'X'
    # elif training_provider_declaration == training_provider_options[2]:
    #     p67 = 'X'
    # elif training_provider_declaration == training_provider_options[3]:
    #     p68 = 'X'
    # elif training_provider_declaration == training_provider_options[4]:
    #     p69 = 'X'
    # elif training_provider_declaration == training_provider_options[5]:
    #     p70 = 'X'
    # elif training_provider_declaration == training_provider_options[6]:
    #     p71 = 'X'
    # elif training_provider_declaration == training_provider_options[7]:
    #     p72 = 'X'

    # justification = st.text_area(
    #         'If there is a discrepancy between Participant self declaration and the PLR, please record justification for level to be reported'
    #     )

    # st.subheader('Does the participant have Basic Skills?')

    # english_options = ['none', 'Entry Level', 'Level 1', 'Level 2+']

    # english_skill = st.selectbox('English', english_options)

    # p74 = '-'
    # p75 = '-'
    # p76 = '-'
    # p77 = '-'

    # if english_skill == english_options[0]:
    #     p74 = 'X'
    # elif english_skill == english_options[1]:
    #     p75 = 'X'
    # elif english_skill == english_options[2]:
    #     p76 = 'X'
    # elif english_skill == english_options[3]:
    #     p77 = 'X'

    # maths_options = ['none', 'Entry Level', 'Level 1', 'Level 2+']

    # maths_skill = st.selectbox('Maths', maths_options)

    # p78 = '-'
    # p79 = '-'
    # p80 = '-'
    # p81 = '-'

    # if maths_skill == maths_options[0]:
    #     p78 = 'X'
    # elif maths_skill == maths_options[1]:
    #     p79 = 'X'
    # elif maths_skill == maths_options[2]:
    #     p80 = 'X'
    # elif maths_skill == maths_options[3]:
    #     p81 = 'X'

    # esol_options = ['none', 'Entry Level', 'Level 1', 'Level 2+']

    # esol_skill = st.selectbox('ESOL', esol_options)

    # p82 = '-'
    # p83 = '-'
    # p84 = '-'
    # p85 = '-'

    # if esol_skill == esol_options[0]:
    #     p82 = 'X'
    # elif esol_skill == esol_options[1]:
    #     p83 = 'X'
    # elif esol_skill == esol_options[2]:
    #     p84 = 'X'
    # elif esol_skill == esol_options[3]:
    #     p85 = 'X'

    # st.subheader('Basic Skills Initial Assessment')
    # st.text(
    #     "Initial Assessment Outcomes – record the levels achieved by the Participant"
    # )

    # maths_options = ['-', 'E1', 'E2', 'E3', '1', '2']

    # maths_level = st.selectbox('Maths Level', maths_options)

    # p86 = ''

    # if maths_level in maths_options[1:]:
    #     p86 = maths_level

    # english_options = ['-', 'E1', 'E2', 'E3', '1', '2']

    # english_level = st.selectbox('English Level', english_options)

    # p87 = ''

    # if english_level in english_options[1:]:
    #     p87 = english_level

    # st.subheader('Numeracy and Literacy Programmes')
    # completion_programmes = st.radio(
    #     'Will the Participant be completing relevant Numeracy and/or Literacy programmes within their learning plan?',
    #     ['Yes', 'No'])
    # p88 = '-'
    # p89 = '-'

    # if completion_programmes == 'Yes':
    #     p88 = 'Y'
    #     p89 = '-'
    # elif completion_programmes == 'No':
    #     p88 = '-'
    #     p89 = 'N'

    # st.subheader('Additional Learning Support')
    # additional_support = st.radio(
    #     'Does the Participant require additional learning and/or learner support?',
    #     ['Yes', 'No'])
    # p90 = '-'
    # p91 = '-'
    # support_details = '-'

    # if additional_support == 'Yes':
    #     p90 = 'Y'
    #     p91 = '-'
    #     support_details = st.text_area(
    #         'If answered \'Yes\' above, please detail how the participant will be supported'
    #     )
    # elif additional_support == 'No':
    #     p90 = '-'
    #     p91 = 'N'

    st.header('Current Skills, Experience, and IAG')

    st.subheader('Highest Level of Education at start')
    education_options = [
        'ISCED 0 - Lacking Foundation skills (below Primary Education)',
        'ISCED 1 - Primary Education',
        'ISCED 2 - GCSE D-G or 3-1/BTEC Level 1/Functional Skills Level 1',
        'ISCED 3 - GCSE A-C or 9-4/AS or A Level/NVQ or BTEC Level 2 or 3',
        'ISCED 4 - N/A',
        'ISCED 5 to 8 - BTEC Level 5 or NVQ Level 4, Foundation Degree, BA, MA or equivalent'
    ]

    # Change from selectbox to multiselect
    selected_levels = st.selectbox(
        'Select the highest level of education at start', education_options)

    # mandatory field validation
    if len(selected_levels)==0:
        # mandatory_fields.extend(['p303'])
        pass

    # Initialize marks
    p93, p94, p95, p96, p97, p98 = '-', '-', '-', '-', '-', '-'

    # Mark selected options
    if education_options[0] in selected_levels:
        p93 = 'X'
    if education_options[1] in selected_levels:
        p94 = 'X'
    if education_options[2] in selected_levels:
        p95 = 'X'
    if education_options[3] in selected_levels:
        p96 = 'X'
    if education_options[4] in selected_levels:
        p97 = 'X'
    if education_options[5] in selected_levels:
        p98 = 'X'

    st.header('Other Information')

    job_role_activities='No job.'
    current_job = st.radio(
    'Are you currently doing job?',
    ['No', 'Yes'])
    if current_job=='Yes':
        st.subheader('Current Job Role and Day to Day Activities')
        job_role_activities = st.text_area(
            'What is your current job role and what are your day to day activities?'
        )


    st.subheader('Career Aspirations')
    career_aspirations = st.text_area('What are your career aspirations? (Please provide details.)')

    training_qualifications_needed='    '
    # st.subheader('Training/Qualifications Needed')
    # training_qualifications_needed = st.text_area(
    #     'What training/qualifications do you need to progress further in your career? (Planned and future training)'
    # )

    barriers_to_achieving_aspirations='    '
    # st.subheader('Barriers to Achieving Career Aspirations')
    # barriers_to_achieving_aspirations = st.text_area(
    #     'What are the barriers to achieving your career aspirations and goals?'
    # )

    # mandatory_fields.extend([f'p{i}' for i in range(99, 103)])

    # st.subheader('Courses/Programs Available')
    # courses_programs_available = st.text_area(
    #     'What courses/programs/activity are available to you in order to meet your and your employer\'s needs?'
    # )

    # st.header('Induction Checklist')


    # funded_by_mayor_of_london = st.checkbox(
    #     'This programme is funded by the Mayor of London')
    # lls_completed = st.checkbox(
    #     'The London Learning Survey (LLS) has been completed and submitted')
    # equality_diversity_policy = st.checkbox(
    #     'Equality and Diversity Policy/Procedure and point of contact')
    # health_safety_policy = st.checkbox(
    #     'Health and Safety Policy/Procedure and point of contact')
    # safeguarding_policy = st.checkbox(
    #     'Safeguarding Policy/Procedure and point of contact')
    # prevent_policy = st.checkbox(
    #     'PREVENT and point of contact (including British Values)')
    # disciplinary_policy = st.checkbox(
    #     'Disciplinary, Appeal and Grievance Policy/Procedures')
    # plagiarism_policy = st.checkbox('Plagiarism, Cheating Policy/Procedure')
    # terms_conditions = st.checkbox(
    #     'Terms and Conditions of Learning and programme content & programme delivery'
    # )



    # Privacy Notice Text
    privacy_notice = """
    Privacy Notice

    This privacy notice is issued by the Education and Skills Funding Agency (ESFA) on behalf of the Secretary of State for the Department of Education (DfE) to inform learners about the Individualised Learner Record (ILR) and how their personal information is used in the ILR. Your personal information is used by the DfE to exercise our functions under article 6(1)(e) of the UK GDPR and to meet our statutory responsibilities, including under the Apprenticeships, Skills, Children and Learning Act 2009.

    The ILR collects data about learners and learning undertaken. Publicly funded colleges, training organisations, local authorities, and employers (FE providers) must collect and return the data to the ESFA each year under the terms of a funding agreement, contract or grant agreement. It helps ensure that public money distributed through the ESFA is being spent in line with government targets. It is also used for education, training, employment, and wellbeing purposes, including research. We retain ILR learner data for 3 years for operational purposes and 66 years for research purposes. For more information about the ILR and the data collected, please see the ILR specification at https://www.gov.uk/government/collections/individualised-learner-record-ilr

    ILR data is shared with third parties where it complies with DfE data sharing procedures and where the law allows it. The DfE and the English European Social Fund (ESF) Managing Authority (or agents acting on their behalf) may contact learners to carry out research and evaluation to inform the effectiveness of training. In these cases, it is part of our statutory duties and we do not need your consent.

    For more information about how your personal data is used and your individual rights, please see the DfE Roles and Responsibilities Personal Information Charter(https://www.gov.uk/government/organisations/department-for-education/about/personal-information-charter) and the ESFA Privacy Notice (https://www.gov.uk/government/publications/esfa-privacy-notice).

    If you would like to get in touch with us, you can contact the DfE in the following ways:
    - Using our online contact form at https://www.gov.uk/government/organisations/department-for-education/about/personal-information-charter.
    - By telephoning the DfE Helpline on 0370 000 2288 or in writing to - Data Protection Officer, Ministerial and Public Communications Division, Department for Education, Piccadilly Gate, Store Street, Manchester, M1 2WD.

    By completing the 'Learner Declaration'. This means that:

    - You understand this provision is delivered by Prevista Ltd or by the named subcontractor on page 1 on behalf Prevista Ltd (or sub-contractor where indicated).
    - You will be the provider know of any changes in your personal circumstances.
    - You fully agree that the provider can process information about you.
    - You understand that the form will be kept until 31st December 2030 at the latest.

    Prevista Ltd will:

    - Provide appropriate guidance and support to the Subcontractor to ensure that they deliver high-quality services.
    - Monitor and evaluate the performance of the Subcontractor regularly to ensure that they meet the agreed-upon standards.
    - Provide the necessary resources and information to the Subcontractor to enable them to carry out their work effectively.
    - Ensure that the Subcontractor complies with all relevant laws and regulations.

    The Subcontractor will:

    - Deliver the agreed-upon services to a high standard and in a timely manner.
    - Comply with all relevant laws and regulations, including health and safety requirements.
    - Provide regular progress reports and updates to Prevista to ensure that they are kept informed of the work being carried out.
    - Work collaboratively with Prevista to ensure that the needs of students and other stakeholders are met.
    """


    # Privacy and Data Protection Information Section
    st.header('Privacy and Data Protection Information')

    # Display the privacy notice text as plain text
    st.text(privacy_notice)


    # Contact preferences
    st.write("Choose Y or N for any of the following if you AGREE to be contacted about courses/learning opportunities")
    contact_surveys = st.radio("For surveys & research", options=["Y", "N"])
    contact_phone = st.radio("Phone", options=["Y", "N"])
    contact_email = st.radio("Email", options=["Y", "N"])
    contact_post = st.radio("Post", options=["Y", "N"])
    # Initialize variables for contact preferences
    contact_surveys_val, contact_phone_val, contact_email_val, contact_post_val = '', '', '', ''
    # Update the variables based on user selections
    contact_surveys_val = 'Y' if contact_surveys == "Y" else 'N'
    contact_phone_val = 'X' if contact_phone == "Y" else 'N'
    contact_email_val = 'X' if contact_email == "Y" else 'N'
    contact_post_val = 'X' if contact_post == "Y" else 'N'

    st.header('Declarations')

    # st.subheader('Provider Confirmation')
    st.text(
        'I hereby confirm that I have read, understood and agree with the contents of this document and above privacy notice, and understand that the programme is funded by the Mayor of London.'
    )


    st.subheader('Participant Declaration')
    st.text_area(
        'Participant Declaration',
        'I certify that I have provided all of the necessary information to confirm my eligibility for the Funded Provision.'
    )


    st.subheader('Participant Signature')

    st.text("Signature:")
    participant_signature = st_canvas(
        fill_color="rgba(255, 255, 255, 1)",  
        stroke_width=5,
        stroke_color="rgb(0, 0, 0)",  # Black stroke color
        background_color="white",  # White background color
        width=400,
        height=150,
        drawing_mode="freedraw",
        key="canvas",
    )

    date_signed = st.date_input(
    label="Date",
    value=date.today(),  # Default date
    min_value=date(1900, 1, 1),  # Minimum selectable date
    max_value=date(2025, 12, 31),  # Maximum selectable date
    help="Choose a date",  # Tooltip text
    format='DD/MM/YYYY'
)
    date_signed = date_signed.strftime("%d-%m-%Y")
    

# ####################################################################################################################################

    submit_button = st.button('Submit')
    if submit_button:
        placeholder_values = {
            'p110': title_mr,
            'p111': title_mrs,
            'p112': title_miss,
            'p113': title_ms,

            'p1': first_name,
            'p2': middle_name,
            'p3': family_name,

            'p114': gender_m,
            'p115': gender_f,
            'p116': other_gender,
            'p117': other_gender_text,

            'p4': date_of_birth,

            'p118': current_age,
            'p119': ethnicity_vars['ethnicity_31'],
            'p120': ethnicity_vars['ethnicity_32'],
            'p121': ethnicity_vars['ethnicity_33'],
            'p122': ethnicity_vars['ethnicity_34'],
            'p123': ethnicity_vars['ethnicity_35'],
            'p124': ethnicity_vars['ethnicity_36'],
            'p125': ethnicity_vars['ethnicity_37'],
            'p126': ethnicity_vars['ethnicity_38'],
            'p127': ethnicity_vars['ethnicity_39'],
            'p128': ethnicity_vars['ethnicity_40'],
            'p129': ethnicity_vars['ethnicity_41'],
            'p130': ethnicity_vars['ethnicity_42'],
            'p131': ethnicity_vars['ethnicity_43'],
            'p132': ethnicity_vars['ethnicity_44'],
            'p133': ethnicity_vars['ethnicity_45'],
            'p134': ethnicity_vars['ethnicity_46'],
            'p135': ethnicity_vars['ethnicity_47'],
            'p136': ethnicity_48,
            'p137': national_insurance_number,
            'p138': house_no_name_street,
            'p139': suburb_village,
            'p140': town_city,
            'p141': county,
            'p142': country_of_domicile,
            'p143': current_postcode,
            'p144': postcode_prior_enrollment,
            'p145': email_address,
            'p146': primary_telephone_number,
            'p147': secondary_telephone_number,
            'p148': next_of_kin,
            'p149': emergency_contact_phone_number,
            'p150': no_member_employed_with_children,
            'p151': no_member_employed_without_children,
            'p152': single_adult_household_with_children,
            'p153': unemployed_single_adult_household,
            'p154': none_of_the_above,            
            'p155': has_disability,
            'p156': no_disability,
            'p157a': vision_impairment_primary,
            'p157b': vision_impairment_secondary,
            'p157c': vision_impairment_tertiary,
            'p158a': hearing_impairment_primary,
            'p158b': hearing_impairment_secondary,
            'p158c': hearing_impairment_tertiary,
            'p159a': mobility_impairment_primary,
            'p159b': mobility_impairment_secondary,
            'p159c': mobility_impairment_tertiary,
            'p160a': complex_disabilities_primary,
            'p160b': complex_disabilities_secondary,
            'p160c': complex_disabilities_tertiary,
            'p161a': social_emotional_difficulties_primary,
            'p161b': social_emotional_difficulties_secondary,
            'p161c': social_emotional_difficulties_tertiary,
            'p162a': mental_health_difficulty_primary,
            'p162b': mental_health_difficulty_secondary,
            'p162c': mental_health_difficulty_tertiary,
            'p163a': moderate_learning_difficulty_primary,
            'p163b': moderate_learning_difficulty_secondary,
            'p163c': moderate_learning_difficulty_tertiary,
            'p164a': severe_learning_difficulty_primary,
            'p164b': severe_learning_difficulty_secondary,
            'p164c': severe_learning_difficulty_tertiary,
            'p165a': dyslexia_primary,
            'p165b': dyslexia_secondary,
            'p165c': dyslexia_tertiary,
            'p166a': dyscalculia_primary,
            'p166b': dyscalculia_secondary,
            'p166c': dyscalculia_tertiary,
            'p167a': autism_spectrum_primary,
            'p167b': autism_spectrum_secondary,
            'p167c': autism_spectrum_tertiary,
            'p168a': aspergers_primary,
            'p168b': aspergers_secondary,
            'p168c': aspergers_tertiary,
            'p169a': temporary_disability_primary,
            'p169b': temporary_disability_secondary,
            'p169c': temporary_disability_tertiary,
            'p170a': speech_communication_needs_primary,
            'p170b': speech_communication_needs_secondary,
            'p170c': speech_communication_needs_tertiary,
            'p171a': physical_disability_primary,
            'p171b': physical_disability_secondary,
            'p171c': physical_disability_tertiary,
            'p172a': specific_learning_difficulty_primary,
            'p172b': specific_learning_difficulty_secondary,
            'p172c': specific_learning_difficulty_tertiary,
            'p173a': medical_condition_primary,
            'p173b': medical_condition_secondary,
            'p173c': medical_condition_tertiary,
            'p174a': other_learning_difficulty_primary,
            'p174b': other_learning_difficulty_secondary,
            'p174c': other_learning_difficulty_tertiary,
            'p175a': other_disability_primary,
            'p175b': other_disability_secondary,
            'p175c': other_disability_tertiary,
            'p176': prefer_not_to_say,
            'p177': additional_info,
            'p178': ex_offender_y,
            'p179': ex_offender_n,
            'p180': ex_offender_choose_not_to_say,

            'p189':homeless_y, 
            'p190':homeless_n,
            'p191':homeless_choose_not_to_say,

            'p181': internally_sourced_val,
            'p182': recommendation_val,
            'p183': event_val,
            'p184': self_referral_val,
            'p185': family_friends_val,
            'p186': other_val,
            'p187': website_val,
            'p188': promotional_material_val,
            'p188a': jobcentre_plus_val,

            'p192': unemployed_val,
            'p193': economically_inactive_val,
            'p194': employed_val,
            'p195': up_to_12_months_val,
            'p196': twelve_months_or_longer_val,
            'p197': jcp_dwp_val,
            'p198': careers_service_val,
            'p199': third_party_val,
            'p200': other_evidence_val,
            'p201': inactive_status_val,
            'p202': inactive_evidence_type_val,
            'p203': inactive_evidence_date_val,  
            'p204': employer_name_val,
            'p205': employer_address_1_val,
            'p206': employer_address_2_val,
            'p207': employer_address_3_val,
            'p208': employer_postcode_val,
            'p209': employer_contact_name_val,
            'p210': employer_contact_position_val,
            'p211': employer_contact_email_val,
            'p212': employer_contact_phone_val,
            'p213': employer_edrs_number_val,
            'p214': living_wage_val,
            'p215a': employment_hours_val_0,
            'p215b': employment_hours_val_6,
            'p216': claiming_benefits_val,
            'p217': sole_claimant_val,
            'p218': universal_credit_val,
            'p219': job_seekers_allowance_val,
            'p220': employment_support_allowance_val,
            'p221': incapacity_benefit_val,
            'p222': personal_independence_payment_val,
            'p223': other_benefit_val,
            'p224': benefit_claim_date_val,                   
            'p225': contact_surveys_val,
            'p226': contact_phone_val,
            'p227': contact_email_val,
            'p228': contact_post_val,

            'p5': nationality,
            'p6': full_uk_passport,
            'p7': full_eu_passport,
            'p8': national_identity_card,
            'p9': hold_settled_status,
            'p10': hold_pre_settled_status,
            'p11': hold_leave_to_remain,
            'p12': not_nationality,
            'p13': passport_non_eu,
            'p14': letter_uk_immigration,
            'p15': passport_endorsed,
            'p16': identity_card,
            'p17': country_of_issue,
            'p18': id_document_reference_number,
            'p19': e01_date_of_issue,
            'p20': e01_date_of_expiry,
            'p21': e01_additional_notes,
            'p22': full_passport_eu,
            'p23': national_id_card_eu,
            'p24': firearms_certificate,
            'p25': birth_adoption_certificate,
            'p26': e02_drivers_license,
            'p27': edu_institution_letter,
            'p28': e02_employment_contract,
            'p29': state_benefits_letter,
            'p30': pension_statement,
            'p31':  northern_ireland_voters_card,
            'p32': e02_other_evidence_text,
            'p33': e02_date_of_issue,
            'p34': e03_drivers_license,
            'p35': bank_statement,
            'p36': pension_statement,
            'p37': mortgage_statement,
            'p38': utility_bill,
            'p39': council_tax_statement,
            'p40': electoral_role_evidence,
            'p41': homeowner_letter,
            'p42': e03_date_of_issue,
            'p43': e03_other_evidence_text,
            'p44': latest_payslip,
            'p45': e04_employment_contract,
            'p46': confirmation_from_employer,
            'p47': redundancy_notice,
            'p48': sa302_declaration,
            'p49': ni_contributions,
            'p50': business_records,
            'p51': companies_house_records,
            'p52': other_evidence_employed,
            'p53': unemployed,
            'p54': e04_date_of_issue,
            'p55': qualification_or_training_y,
            'p56': qualification_or_training_n,
            'p57': course_details + ' ' + funding_details,
            'p58': p58,
            'p59': p59,
            'p60': p60,
            'p61': p61,
            'p62': p62,
            'p63': p63,
            'p64': p64,

            'p60z' : p60z,
            'p60a' : p60a,
            'p61z' : p61z,
            'p61a' : p61a,
            'p63z' : p63z,
            'p63a' : p63a,
            'p63b' : p63b,


            # 'p65': p65,
            # 'p66': p66,
            # 'p67': p67,
            # 'p68': p68,
            # 'p69': p69,
            # 'p70': p70,
            # 'p71': p71,
            # 'p72': p72,
            # 'p73': justification,
            # 'p74': p74,
            # 'p75': p75,
            # 'p76': p76,
            # 'p77': p77,
            # 'p78': p78,
            # 'p79': p79,
            # 'p80': p80,
            # 'p81': p81,
            # 'p82': p82,
            # 'p83': p83,
            # 'p84': p84,
            # 'p85': p85,
            # 'p86': p86,
            # 'p87': p87,
            # 'p88': p88,
            # 'p89': p89,
            # 'p90': p90,
            # 'p91': p91,
            # 'p92': support_details,
            'p93': p93,
            'p94': p94,
            'p95': p95,
            'p96': p96,
            'p97': p97,
            'p98': p98,
            'p99': job_role_activities,
            'p100': career_aspirations,
            'p101': training_qualifications_needed,
            'p102': barriers_to_achieving_aspirations,
            # 'p103': courses_programs_available,
            # 'p113': participant_signature,
            'p231': date_signed,
            
            # for validation
            'p300': household_filled,
            'p301': e02_filled,
            'p302': e03_filled,
            'p303': len(selected_levels),
            'p304': referrall,
            'p305': specify_refereel,
            

        }
        
        # progress_bar(5)

        # mandatory fields validation
        
        # exclude_fields = {'p1000', 'p1', 'p2', 'p3', 'p5', 'p7', 'p8', 'p10', 'p11', 'p12', 'p13', 'p15', 'p16', 'p17', 'p18', 'p32', 'p43', 'p73', 'p86', 'p87', 'p92', 'p99', 'p100', 'p101', 'p102', 'p103', 'p9', 'p14', 'p19', 'p20', 'p21', 'p111', 'p112', 'p113', 'p115', 'p116', 'p117', 'p119', 'p120', 'p121', 'p122', 'p123', 'p124', 'p125', 'p126', 'p127', 'p128', 'p129', 'p130', 'p131', 'p132', 'p133', 'p134', 'p135', 'p137', 'p138', 'p139', 'p140', 'p141', 'p142', 'p143', 'p144', 'p145', 'p146', 'p147', 'p148', 'p149', 'p150'}     # exclude fields
        
        mandatory_fields.extend([f'p{i}' for i in range(0, 0)])

        # Remove excluded fields from mandatory_fields
        mandatory_fields = [field for field in mandatory_fields if field not in exclude_fields]

        missing_fields = validate_inputs(placeholder_values, mandatory_fields)  # get the list of missing mandatory inputs
        if missing_fields:
            st.warning(f"Please fill out all the fields.")
            st.text(f'Error Code: {missing_fields}')
            # st.text('LENGTH:', len(specify_refereel))
            
        else:   
            # Define input and output paths
            template_file = "ph_gla_v3.docx"
            modified_file = f"GLA_Form_Submission_{first_name}_{middle_name}_{family_name}.docx"

            if len(participant_signature.json_data['objects']) != 0:
                # Convert the drawing to a PIL image and save it
                signature_path = 'signature_image.png'
                signature_image = PILImage.fromarray(
                    participant_signature.image_data.astype('uint8'), 'RGBA')
                signature_image.save(signature_path)

                replace_placeholders(template_file, modified_file, placeholder_values, signature_path)

                # Email

                # Sender email credentials

                # Credentials: Streamlit host st.secrets
                sender_email = st.secrets["sender_email"]
                sender_password = st.secrets["sender_password"]

                # Credentials: Local env
                # load_dotenv()                                     # uncomment import of this library!
                # sender_email = os.getenv('EMAIL')
                # sender_password = os.getenv('PASSWORD')

                receiver_email = sender_email
                # receiver_email = 'mohamedr@prevista.co.uk'
                
                subject = f"GLA: {first_name} {middle_name} {family_name} {date.today()} {specify_refereel}"

                body = "GLA Form submitted. Please find attached files."

                # Local file path
                local_file_path = modified_file

                # Send email with attachments
                if files or local_file_path:
                    send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files, local_file_path)
                    st.success("Response sent successfully!")
                else:
                    st.warning("Please upload at least one file or specify a local file.")
        
            else:
                st.warning("Please draw your signature.")


#  Custome Functions 
# ####################################################################################################################################
def validate_inputs(inputs, mandatory_fields):
    """Check if all mandatory input fields are filled and return the list of missing fields."""
    missing_fields = []
    for key, value in inputs.items():
        if key in mandatory_fields and (value is None or value == '' or value == 0):
            missing_fields.append(key)
    return missing_fields

def resize_image_to_fit_cell(image, max_width, max_height):
    width, height = image.size
    aspect_ratio = width / height

    if width > max_width:
        width = max_width
        height = int(width / aspect_ratio)

    if height > max_height:
        height = max_height
        width = int(height * aspect_ratio)

    return image.resize((width, height))


def replace_placeholders(template_file, modified_file, placeholder_values, signature_path):
    try:
        print(f"Copying template file '{template_file}' to '{modified_file}'...")
        shutil.copy(template_file, modified_file)

        print(f"Opening document '{modified_file}'...")
        doc = Document(modified_file)

        # Function to convert value to string, handling datetime.date objects
        def convert_to_str(value):
            if isinstance(value, date):
                return value.strftime('%Y-%m-%d')  # Convert date to string
            return str(value)  # Convert other types to string

        # Compile regular expressions for all placeholders
        placeholders = {re.escape(key): convert_to_str(value) for key, value in placeholder_values.items()}
        placeholders_pattern = re.compile(r'\b(' + '|'.join(placeholders.keys()) + r')\b')

        # Replace placeholders in paragraphs
        print("Replacing placeholders in paragraphs...")
        for para in doc.paragraphs:
            original_text = para.text
            updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
            if original_text != updated_text:
                print(f"Updated paragraph text: '{original_text}' -> '{updated_text}'")
                para.text = updated_text

        # Replace placeholders in tables
        print("Replacing placeholders in tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], para.text)
                        if original_text != updated_text:
                            print(f"Updated table cell text: '{original_text}' -> '{updated_text}'")
                            para.text = updated_text

                    # Inspect cell runs
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run_text = run.text
                            run_updated_text = placeholders_pattern.sub(lambda match: placeholders[re.escape(match.group(0))], run_text)
                            if run_text != run_updated_text:
                                print(f"Updated run text in table cell: '{run_text}' -> '{run_updated_text}'")
                                run.text = run_updated_text

        # Check and handle signature placeholder
        print("Inspecting document for 'p230' placeholder...")
        signature_placeholder_found = False

        # Check paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()  # Remove any extra spaces around text
            while 'p230' in para_text:
                print(f"Found 'p230' in paragraph: '{para_text}'")
                para_text = para_text.replace('p230', '').strip()  # Remove 'p230' and any leading/trailing spaces
                para.text = para_text
                resized_image_path = 'resized_signature_image.png'
                
                try:
                    # Open and resize the image
                    print(f"Opening image file: {signature_path}")
                    resized_image = PILImage.open(signature_path)
                    print(f"Original image size: {resized_image.size}")
                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                    resized_image.save(resized_image_path)  # Save resized image to a file
                    print(f"Resized image saved to: {resized_image_path}")
                    
                    # Add picture to the paragraph
                    print(f"Adding picture to paragraph from path: {resized_image_path}")
                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                    print("Inserted signature image into paragraph.")
                    signature_placeholder_found = True
                except Exception as img_e:
                    print(f"An error occurred with image processing: {img_e}")

        # Check table cells again in case the placeholder was missed
        if not signature_placeholder_found:
            print("Checking table cells for 'p230'...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para_text = para.text.strip()
                            while 'p230' in para_text:
                                print(f"Found 'p230' in table cell paragraph: '{para_text}'")
                                para_text = para_text.replace('p230', '').strip()
                                para.text = para_text
                                resized_image_path = 'resized_signature_image.png'
                                
                                try:
                                    # Open and resize the image
                                    print(f"Opening image file: {signature_path}")
                                    resized_image = PILImage.open(signature_path)
                                    print(f"Original image size: {resized_image.size}")
                                    resized_image = resize_image_to_fit_cell(resized_image, 200, 50)
                                    resized_image.save(resized_image_path)  # Save resized image to a file
                                    print(f"Resized image saved to: {resized_image_path}")
                                    
                                    # Add picture to the table cell
                                    print(f"Adding picture to table cell from path: {resized_image_path}")
                                    para.add_run().add_picture(resized_image_path, width=Inches(2))
                                    print("Inserted signature image into table cell.")
                                    signature_placeholder_found = True
                                except Exception as img_e:
                                    print(f"An error occurred with image processing: {img_e}")

        if not signature_placeholder_found:
            print("No signature placeholder found.")

        # Save the modified document
        print(f"Saving modified document '{modified_file}'...")
        doc.save(modified_file)
        print(f"Document modification complete: '{modified_file}'")

    except Exception as e:
        print(f"An error occurred: {e}")

    # file download button
    with open(modified_file, 'rb') as f:
        file_contents = f.read()
        st.download_button(
            label="Download Your Response",
            data=file_contents,
            file_name=modified_file,
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )



# Function to send email with attachments (Handle Local + Uploaded)
def send_email_with_attachments(sender_email, sender_password, receiver_email, subject, body, files, local_file_path=None):
    msg = EmailMessage()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject
    msg.set_content(body)

    # Attach uploaded files
    for uploaded_file in files:
        uploaded_file.seek(0)  # Move to the beginning of the UploadedFile
        msg.add_attachment(uploaded_file.read(), maintype='application', subtype='octet-stream', filename=uploaded_file.name)

    # Attach local file if specified
    if local_file_path:
        with open(local_file_path, 'rb') as f:
            file_data = f.read()
            file_name = local_file_path.split('/')[-1]
            msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=file_name)

    # Use the SMTP server for sending the email
    with smtplib.SMTP('smtp.office365.com', 587) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)


# Function to add a checkbox with a file upload option
def add_checkbox_with_upload(label, key_prefix):
    global files
    checked = st.checkbox(label, key=f"{key_prefix}_checkbox")
    if checked:
        st.text(f'Please upload a copy of your {label}')
        uploaded_file = st.file_uploader(f"Upload {label}", type=['pdf', 'jpg', 'jpeg', 'png', 'docx'], key=f"{key_prefix}_uploader")
        if uploaded_file is not None:
            files.append(uploaded_file)
        uploaded_file_1 = st.file_uploader(f"Optional - Upload Back Side of The Document", type=['pdf', 'jpg', 'jpeg', 'png', 'docx'], key=f"{key_prefix}_uploader_1")
        if uploaded_file_1 is not None:
            files.append(uploaded_file_1)
        return 'X'
    else:
        return '-'

# Function to handle file upload
def handle_file_upload(label, key_prefix):
    global files
    st.text(f'Please upload a copy of your {label}')
    uploaded_file = st.file_uploader(f"Upload {label}", type=['pdf', 'jpg', 'jpeg', 'png', 'docx'], key=key_prefix)
    if uploaded_file is not None:
        files.append(uploaded_file)
        return 'X'
    else:
        return '-'
        
def calculate_age(born):
    today = date.today()
    return today.year - born.year - ((today.month, today.day) < (born.month, born.day))

    st.markdown(scroll_script, unsafe_allow_html=True)

def progress_bar(duration_seconds):
    """Displays a progress bar that fills over the specified duration."""
    progress_bar = st.progress(0)
    
    # Number of updates per second for smoother progress
    updates_per_second = 20
    # Time to wait between updates
    sleep_time = 1 / updates_per_second
    # Total number of updates
    total_updates = duration_seconds * updates_per_second
    
    for i in range(total_updates + 1):
        # Update the progress bar
        progress = i / total_updates
        progress_bar.progress(progress)
        # Sleep for the calculated time
        time.sleep(sleep_time)
    # st.write("Progress complete!")

if __name__ == '__main__':
    app()
# streamlit run app.py --server.port 8501
# Dev : https://linkedin.com/in/osamatech786